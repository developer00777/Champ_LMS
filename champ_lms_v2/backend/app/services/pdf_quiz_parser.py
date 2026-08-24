"""
Turn a question-and-answer document (PDF or Word .docx) into structured
TestQuestion data.

Text extraction is format-specific; everything after it is shared. There is no
standard format for a "quiz document", so this runs a deterministic pattern
parser first (free, instant, handles the common numbered-question +
lettered-options + answer-key layouts) and leaves the AI parser in
ai_service.parse_questions_from_text() as the fallback for odd layouts.

Whatever comes out is treated as a DRAFT: the admin reviews and fixes it in the
UI before the test can be published. That's deliberate — silently publishing a
mis-parsed exam is worse than making someone glance at it.
"""
from __future__ import annotations

import io
import re

from app.models.test_series import (
    QUESTION_TYPE_MCQ,
    QUESTION_TYPE_WRITTEN,
    TestQuestion,
)

# --- Question stems: "1." / "1)" / "Q1." / "Q.1" / "Question 1:" -------------
_Q_START = re.compile(
    r"^\s*(?:Q(?:uestion)?\s*\.?\s*)?(\d{1,3})\s*[.)\]:-]\s+(?P<text>\S.*)$",
    re.IGNORECASE,
)
# A bare "Q." with no number still starts a question.
_Q_START_BARE = re.compile(r"^\s*Q(?:uestion)?\s*[.):]\s*(?P<text>\S.*)$", re.IGNORECASE)

# --- Options: "A)" / "(A)" / "A." / "a -" -----------------------------------
_OPTION = re.compile(
    r"^\s*\(?\s*([A-Ha-h])\s*[).\]:-]\s*(?P<text>\S.*)$",
)
# Numeric options ("1) ... 2) ...") only inside an options block, since they'd
# otherwise be indistinguishable from question numbers.
_OPTION_NUM = re.compile(r"^\s*\(?\s*([1-8])\s*[).\]:-]\s*(?P<text>\S.*)$")

# --- Inline answer: "Answer: B" / "Ans - 2" / "Correct Answer : (c)" --------
_INLINE_ANSWER = re.compile(
    r"^\s*(?:correct\s*)?ans(?:wer)?\s*(?:key)?\s*[:.\-)]?\s*\(?\s*([A-Ha-h]|[1-8])\s*\)?\s*\.?\s*$",
    re.IGNORECASE,
)
# Answer plus explanation on one line: "Answer: B — because ..."
_INLINE_ANSWER_EXPL = re.compile(
    r"^\s*(?:correct\s*)?ans(?:wer)?\s*(?:key)?\s*[:.\-)]?\s*\(?\s*([A-Ha-h]|[1-8])\s*\)?\s*[-–—:.]\s*(?P<expl>\S.*)$",
    re.IGNORECASE,
)
_EXPLANATION = re.compile(
    r"^\s*(?:explanation|rationale|why|reason)\s*[:.\-]\s*(?P<text>\S.*)$", re.IGNORECASE
)
_TOPIC = re.compile(
    r"^\s*(?:topic|subject|area|category)\s*[:.\-]\s*(?P<text>\S.*)$", re.IGNORECASE
)
_MARKS = re.compile(
    r"^\s*(?:marks?|points?|score)\s*[:.\-]\s*(\d{1,3})\s*$", re.IGNORECASE
)

# --- Trailing answer key: header, then "1. B" / "1-B" / "1: (b)" pairs ------
_KEY_HEADER = re.compile(
    r"^\s*(?:answer\s*key|answers?|solutions?|answer\s*sheet)\s*[:.]?\s*$", re.IGNORECASE
)
_KEY_PAIR = re.compile(r"(\d{1,3})\s*[).:\-=]?\s*\(?\s*([A-Ha-h])\s*\)?")


# Phrasing that marks an option-less question as a real written-answer prompt
# rather than a heading the numbered-stem pattern caught by mistake.
_WRITTEN_HINT = re.compile(
    r"\b(explain|describe|discuss|justify|elaborate|outline|summari[sz]e|"
    r"compare|contrast|evaluate|analys[ei]|analyz[ei]|define|illustrate|"
    r"comment on|write (?:a|an|short|briefly)|state (?:why|how)|"
    r"in your own words|give (?:reasons|examples)|what (?:steps|would you)|"
    r"how (?:would|does|do) you|why (?:do|does|is|are|would))\b",
    re.IGNORECASE,
)
# An explicit marks/word allowance is also a strong signal of a written answer.
_WRITTEN_ALLOWANCE = re.compile(
    r"\(?\s*(?:in\s+)?(?:about\s+|approx(?:\.|imately)?\s+|max(?:imum)?\s+)?"
    r"(\d{2,4})\s*words?\s*\)?", re.IGNORECASE
)
# A "Answer:" line with prose (not a single letter) is a model answer.
_MODEL_ANSWER = re.compile(
    r"^\s*(?:model\s+|expected\s+|sample\s+|correct\s+)?ans(?:wer)?\s*[:.\-]\s*(?P<text>\S.{4,})$",
    re.IGNORECASE,
)


class PdfParseError(Exception):
    """
    Raised when the upload isn't a readable question document at all.

    Name kept for backwards compatibility with existing importers; it covers
    Word documents too now.
    """


# Formats we can pull text out of, by lowercase file extension.
PDF_EXTENSIONS = (".pdf",)
DOCX_EXTENSIONS = (".docx",)
SUPPORTED_EXTENSIONS = PDF_EXTENSIONS + DOCX_EXTENSIONS

# Magic bytes, so a mislabelled extension or a browser that sends a vague
# content-type still routes to the right extractor.
_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")  # .docx is a zip
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"  # legacy .doc / .xls


def _looks_like_pdf(data: bytes) -> bool:
    # Some producers emit junk before the header, so scan the first block.
    return _PDF_MAGIC in data[:1024]


def _looks_like_docx(data: bytes) -> bool:
    return data[:4] in (m[:4] for m in _ZIP_MAGIC)


def _looks_like_legacy_doc(data: bytes) -> bool:
    return data[:8] == _OLE_MAGIC


def extract_text_from_docx(data: bytes) -> str:
    """
    Pull text out of a Word .docx. Paragraphs and table cells both count — quiz
    papers are very often laid out as a table (one row per question, or a
    question/answer two-column grid), and skipping tables would silently lose
    most of the document.

    Import is local so a missing python-docx surfaces as a clear error on this
    endpoint instead of breaking app startup.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise PdfParseError(
            "Word support unavailable: the 'python-docx' package is not installed."
        ) from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise PdfParseError(
            f"Could not read this Word document: {exc}. If it was saved as an "
            "older .doc file, re-save it as .docx and upload again."
        ) from exc

    lines: list[str] = []

    def emit(text: str) -> None:
        text = text.replace("\xa0", " ").rstrip()
        if text.strip():
            lines.append(text)

    for para in doc.paragraphs:
        emit(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # A row's cells are de-duplicated because merged cells repeat their
            # text across every grid position they span.
            seen: list[str] = []
            for cell in cells:
                if cell and (not seen or cell != seen[-1]):
                    seen.append(cell)
            # Cells often each hold a full line ("A) foo"), so emit them
            # separately rather than joining a row into one unparseable line.
            for cell in seen:
                for cell_line in cell.splitlines():
                    emit(cell_line)

    return "\n".join(lines)


def extract_text(data: bytes) -> str:
    """
    Pull text out of a PDF. Import is local so a missing pypdf surfaces as a
    clear error on this endpoint instead of breaking app startup.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise PdfParseError(
            "PDF support unavailable: the 'pypdf' package is not installed."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise PdfParseError(f"Could not read PDF: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # Many "encrypted" PDFs use an empty owner password and open fine.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise PdfParseError(
                "This PDF is password-protected. Remove the password and re-upload."
            ) from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # * one unreadable page shouldn't sink the whole upload
            pages.append("")
    return "\n".join(pages)


def _letter_to_index(token: str) -> int | None:
    token = token.strip().lower()
    if not token:
        return None
    if token.isdigit():
        n = int(token)
        return n - 1 if n >= 1 else None
    return ord(token) - ord("a")


def _split_answer_key(lines: list[str]) -> tuple[list[str], dict[int, int]]:
    """
    Detect a trailing answer-key section and peel it off.
    Returns (body_lines, {question_number: correct_index}).
    """
    for i, line in enumerate(lines):
        if not _KEY_HEADER.match(line):
            continue
        tail_lines = lines[i + 1 :]
        tail = "\n".join(tail_lines)
        pairs = _KEY_PAIR.findall(tail)
        # Guard against a stray "Answers:" mid-document truncating the paper:
        # accept a single-pair key only when everything after the header is
        # nothing but number/letter pairs (a real key at the end of the file).
        tail_is_only_pairs = bool(tail_lines) and all(
            _KEY_PAIR.fullmatch(ln.strip()) for ln in tail_lines
        )
        if len(pairs) >= 2 or (len(pairs) == 1 and tail_is_only_pairs):
            key: dict[int, int] = {}
            for num, letter in pairs:
                idx = _letter_to_index(letter)
                if idx is not None and idx >= 0:
                    key[int(num)] = idx
            if key:
                return lines[:i], key
    return lines, {}


def _flush(cur: dict) -> TestQuestion | None:
    """Turn the in-progress accumulator into a TestQuestion, if it's usable."""
    text = " ".join(cur["question"]).strip()
    if not text:
        return None
    return TestQuestion(
        question=text,
        options=[o.strip() for o in cur["options"] if o.strip()],
        correct_index=cur["correct_index"],
        explanation=cur["explanation"] or None,
        topic=cur["topic"] or None,
        marks=cur["marks"],
        expected_answer=(cur.get("expected_answer") or None),
    )


def parse_questions(text: str) -> list[TestQuestion]:
    """
    Pattern-parse questions out of extracted PDF text.

    Recognises, per question: a numbered stem, lettered/numbered options, an
    inline answer, an optional explanation/topic/marks, plus a trailing
    answer-key section that back-fills answers by question number.
    """
    raw_lines = [ln.rstrip() for ln in text.splitlines()]
    lines = [ln for ln in raw_lines if ln.strip()]
    body, answer_key = _split_answer_key(lines)

    questions: list[TestQuestion] = []
    numbers: list[int | None] = []  # parallel to `questions`, for answer-key lookup

    cur: dict | None = None
    cur_num: int | None = None
    in_options = False
    numeric_options = False  # this question's options are "1)…2)…", not "A)…B)…"
    in_answer = False  # inside a multi-line written model answer
    last_q_num = 0

    def start(num: int | None, stem: str) -> None:
        nonlocal cur, cur_num, in_options, numeric_options, in_answer, last_q_num
        cur = {
            "question": [stem], "options": [], "correct_index": None,
            "explanation": None, "topic": None, "marks": 1,
            "expected_answer": None,
        }
        cur_num, in_options, numeric_options = num, False, False
        in_answer = False
        if num is not None:
            last_q_num = num

    def commit() -> None:
        nonlocal cur
        if cur:
            q = _flush(cur)
            if q:
                questions.append(q); numbers.append(cur_num)
        cur = None

    for line in body:
        stripped = line.strip()

        # A lettered marker ("A)", "(b)") is unambiguous — always an option.
        opt = _OPTION.match(stripped)

        # A numeric marker ("1)", "2.") is ambiguous — it could be an option or
        # the next question. Resolve it by sequence: options run 1,2,3… within the
        # current question, while question numbers ascend globally. Question wins
        # ties, so a numbered paper without lettered options still parses.
        m_q = None if opt else _Q_START.match(stripped)
        if m_q and cur is not None:
            n = int(m_q.group(1))
            if numeric_options and n == len(cur["options"]) + 1:
                # Mid-run of numeric options — continues the run. Checked before
                # the global-sequence rule, which would otherwise steal option
                # "2)" as question 2 when the stem was "Question 1:".
                opt, m_q = _OPTION_NUM.match(stripped), None
            elif n == last_q_num + 1:
                pass  # next question in sequence — treat as a stem
            elif not cur["options"] and n == 1:
                # first option of a numeric-optioned question
                opt, m_q = _OPTION_NUM.match(stripped), None

        if m_q is None and not in_options:
            bare = _Q_START_BARE.match(stripped)
            if bare:
                commit()
                start(None, bare.group("text"))
                continue

        # A numbered stem starts a new question.
        if m_q:
            commit()
            start(int(m_q.group(1)), m_q.group("text"))
            continue

        if cur is None:
            continue  # preamble before the first question

        if opt:
            cur["options"].append(opt.group("text"))
            in_options = True
            in_answer = False  # options end any model-answer run
            if opt.group(1).isdigit():
                numeric_options = True
            continue

        m_ans_expl = _INLINE_ANSWER_EXPL.match(stripped)
        if m_ans_expl:
            cur["correct_index"] = _letter_to_index(m_ans_expl.group(1))
            cur["explanation"] = m_ans_expl.group("expl").strip()
            continue

        m_ans = _INLINE_ANSWER.match(stripped)
        if m_ans:
            cur["correct_index"] = _letter_to_index(m_ans.group(1))
            continue

        # "Answer: <prose>" on a question with no lettered options is a model
        # answer for a written question, not an MCQ key. Checked after the
        # single-letter patterns above so "Answer: B" is still an MCQ key.
        if not cur["options"]:
            m_model = _MODEL_ANSWER.match(stripped)
            if m_model:
                cur["expected_answer"] = m_model.group("text").strip()
                in_answer = True
                continue

        m_expl = _EXPLANATION.match(stripped)
        if m_expl:
            cur["explanation"] = m_expl.group("text").strip()
            continue

        m_topic = _TOPIC.match(stripped)
        if m_topic:
            cur["topic"] = m_topic.group("text").strip()
            continue

        m_marks = _MARKS.match(stripped)
        if m_marks:
            cur["marks"] = max(1, int(m_marks.group(1)))
            continue

        # Continuation line: extends the model answer, the last option, or the stem.
        if in_answer and cur["expected_answer"]:
            cur["expected_answer"] += " " + stripped
        elif in_options and cur["options"]:
            cur["options"][-1] += " " + stripped
        elif cur["explanation"]:
            cur["explanation"] += " " + stripped
        else:
            cur["question"].append(stripped)

    if cur:
        q = _flush(cur)
        if q:
            questions.append(q); numbers.append(cur_num)

    # Back-fill from the trailing answer key for anything still unanswered.
    if answer_key:
        for i, q in enumerate(questions):
            if q.correct_index is not None:
                continue
            num = numbers[i] if i < len(numbers) else None
            idx = answer_key.get(num) if num is not None else None
            if idx is None:
                idx = answer_key.get(i + 1)  # fall back to positional order
            if idx is not None and idx < len(q.options):
                q.correct_index = idx

    return _classify(questions)


def _looks_written(q: TestQuestion) -> bool:
    """
    Is this option-less question a real written-answer prompt?

    Option-less stems used to be dropped wholesale as heading noise. Written
    questions look identical structurally, so they are told apart by content:
    an instruction verb ("explain", "describe"), a word allowance, a model
    answer captured as the explanation, or simply enough words to be a real
    question rather than a heading.
    """
    text = q.question.strip()
    if _WRITTEN_HINT.search(text):
        return True
    if _WRITTEN_ALLOWANCE.search(text):
        return True
    if q.expected_answer:
        return True
    # A heading is short and rarely a question. Require either a question mark
    # or a sentence of real length before treating it as a prompt.
    return text.endswith("?") and len(text.split()) >= 4


def _classify(questions: list[TestQuestion]) -> list[TestQuestion]:
    """
    Tag each parsed question as mcq or written, and drop leftover noise.

    Anything with 2+ options is an MCQ, exactly as before. Option-less entries
    survive only when they read like a written-answer prompt.
    """
    out: list[TestQuestion] = []
    for q in questions:
        if len(q.options) >= 2:
            q.question_type = QUESTION_TYPE_MCQ
            out.append(q)
            continue
        if _looks_written(q):
            q.question_type = QUESTION_TYPE_WRITTEN
            q.options = []
            q.correct_index = None
            m = _WRITTEN_ALLOWANCE.search(q.question)
            if m:
                try:
                    q.max_words = int(m.group(1))
                except ValueError:
                    pass
            out.append(q)
        # else: a heading or instruction line — dropped, as before.
    return out


def parse_document(data: bytes, filename: str | None = None) -> tuple[list[TestQuestion], str]:
    """
    Extract text from a PDF or Word .docx, then pattern-parse it.
    Returns (questions, extracted_text).

    Format is decided by magic bytes first and the filename only as a
    tie-breaker, so a .docx renamed to .pdf (or an unhelpful content-type from
    the browser) still parses instead of failing confusingly.
    """
    if not data:
        raise PdfParseError("The uploaded file is empty.")

    ext = ""
    if filename and "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if _looks_like_pdf(data):
        text, kind = extract_text(data), "PDF"
    elif _looks_like_docx(data):
        text, kind = extract_text_from_docx(data), "Word document"
    elif _looks_like_legacy_doc(data):
        raise PdfParseError(
            "This is an older Word .doc file, which can't be read directly. "
            "Open it in Word or Google Docs and save it as .docx (or export to "
            "PDF), then upload again."
        )
    elif ext in PDF_EXTENSIONS:
        # Named .pdf but no %PDF header — let pypdf produce the precise error.
        text, kind = extract_text(data), "PDF"
    elif ext in DOCX_EXTENSIONS:
        text, kind = extract_text_from_docx(data), "Word document"
    else:
        raise PdfParseError(
            "Unsupported file type. Upload a PDF or a Word .docx question paper."
        )

    if not text.strip():
        if kind == "PDF":
            raise PdfParseError(
                "No text found in this PDF — it may be a scanned image. "
                "Upload a text-based PDF, or add the questions manually."
            )
        raise PdfParseError(
            "No text found in this Word document. If the questions are inside "
            "images or text boxes, paste them into the document body (or add "
            "them manually) and upload again."
        )
    return parse_questions(text), text


def parse_pdf(data: bytes) -> tuple[list[TestQuestion], str]:
    """Backwards-compatible alias for PDF-only callers."""
    return parse_document(data, filename="upload.pdf")
